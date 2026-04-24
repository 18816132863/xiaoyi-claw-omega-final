from pathlib import Path
import os
#!/usr/bin/env python3

PROJECT_ROOT = Path(__file__).resolve().parents[1]
"""
生成专业格式的完整文档 - 带图表、详细内容
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
import numpy as np
from datetime import datetime

def add_chart_image(doc, chart_path, width_inches=6):
    """添加图表到文档"""
    doc.add_picture(chart_path, width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_health_report_with_charts():
    """创建带图表的健康数据周报"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('健康数据周报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('2026年4月12日 - 2026年4月18日')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 用户信息
    p = doc.add_paragraph()
    p.add_run('用户: ').bold = True
    p.add_run('张三\n')
    p.add_run('年龄: ').bold = True
    p.add_run('30岁\n')
    p.add_run('性别: ').bold = True
    p.add_run('男\n')
    p.add_run('身高: ').bold = True
    p.add_run('175cm\n')
    p.add_run('体重: ').bold = True
    p.add_run('70kg')
    
    doc.add_page_break()
    
    # 第一部分：数据概览
    doc.add_heading('一、数据概览', level=1)
    
    # 创建数据表格
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    headers = ['指标', '本周平均', '上周平均', '变化趋势']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 数据
    data = [
        ('睡眠时长', '7.2小时', '6.8小时', '↑ 5.9%'),
        ('运动时长', '35分钟', '28分钟', '↑ 25%'),
        ('步数', '8,500步', '7,200步', '↑ 18%'),
        ('心率', '72次/分', '75次/分', '↓ 4%')
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # 生成趋势图
    days = ['周一', '周二', '周三', '周四', '周五', '周六', '周日']
    sleep_data = [7.5, 6.8, 7.2, 7.0, 7.8, 7.5, 6.5]
    steps_data = [8500, 7200, 12300, 9100, 8800, 5200, 7800]
    
    # 睡眠趋势图
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(days, sleep_data, marker='o', linewidth=2, markersize=8, color='#2E86AB')
    ax.fill_between(days, sleep_data, alpha=0.3, color='#2E86AB')
    ax.set_ylabel('睡眠时长（小时）', fontsize=12)
    ax.set_title('本周睡眠趋势', fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3)
    ax.set_ylim(5, 9)
    
    # 添加目标线
    ax.axhline(y=7.5, color='green', linestyle='--', label='目标: 7.5小时')
    ax.legend()
    
    plt.tight_layout()
    sleep_chart = 'str(PROJECT_ROOT)/skills/data-tracker/output/sleep_chart.png'
    plt.savefig(sleep_chart, dpi=150, bbox_inches='tight')
    plt.close()
    
    # 步数趋势图
    fig, ax = plt.subplots(figsize=(10, 5))
    colors = ['#FF6B6B' if x < 8000 else '#4ECDC4' for x in steps_data]
    bars = ax.bar(days, steps_data, color=colors, alpha=0.8)
    ax.set_ylabel('步数', fontsize=12)
    ax.set_title('本周步数统计', fontsize=14, fontweight='bold')
    ax.axhline(y=8000, color='green', linestyle='--', label='目标: 8000步')
    ax.legend()
    ax.grid(True, alpha=0.3, axis='y')
    
    # 添加数值标签
    for bar, val in zip(bars, steps_data):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 200, 
                f'{val:,}', ha='center', va='bottom', fontsize=10)
    
    plt.tight_layout()
    steps_chart = 'str(PROJECT_ROOT)/skills/data-tracker/output/steps_chart.png'
    plt.savefig(steps_chart, dpi=150, bbox_inches='tight')
    plt.close()
    
    doc.add_page_break()
    
    # 第二部分：睡眠分析
    doc.add_heading('二、睡眠质量分析', level=1)
    
    add_chart_image(doc, sleep_chart)
    
    doc.add_heading('睡眠质量评分', level=2)
    
    # 睡眠质量详情
    p = doc.add_paragraph()
    p.add_run('平均睡眠时长: ').bold = True
    p.add_run('7.2小时\n')
    p.add_run('深度睡眠占比: ').bold = True
    p.add_run('23%\n')
    p.add_run('浅度睡眠占比: ').bold = True
    p.add_run('52%\n')
    p.add_run('REM睡眠占比: ').bold = True
    p.add_run('25%\n')
    p.add_run('睡眠质量评分: ').bold = True
    p.add_run('85分/100分')
    
    doc.add_heading('睡眠建议', level=2)
    doc.add_paragraph('✅ 睡眠质量良好，继续保持规律作息', style='List Bullet')
    doc.add_paragraph('⚠️ 建议增加深度睡眠时长，可通过睡前放松活动改善', style='List Bullet')
    doc.add_paragraph('💡 建议睡前1小时避免使用电子设备', style='List Bullet')
    doc.add_paragraph('💡 保持卧室温度在18-22°C，湿度在40-60%', style='List Bullet')
    
    doc.add_page_break()
    
    # 第三部分：运动分析
    doc.add_heading('三、运动情况分析', level=1)
    
    add_chart_image(doc, steps_chart)
    
    doc.add_heading('运动统计', level=2)
    
    p = doc.add_paragraph()
    p.add_run('运动天数: ').bold = True
    p.add_run('5天/7天\n')
    p.add_run('运动类型: ').bold = True
    p.add_run('跑步、游泳、力量训练\n')
    p.add_run('总运动时长: ').bold = True
    p.add_run('175分钟\n')
    p.add_run('消耗卡路里: ').bold = True
    p.add_run('2,100千卡\n')
    p.add_run('平均心率: ').bold = True
    p.add_run('128次/分')
    
    doc.add_heading('运动建议', level=2)
    doc.add_paragraph('✅ 运动量达标，每周运动5天表现优秀', style='List Bullet')
    doc.add_paragraph('⚠️ 建议增加力量训练频率，每周至少2次', style='List Bullet')
    doc.add_paragraph('💡 可尝试HIIT训练，提高燃脂效率', style='List Bullet')
    doc.add_paragraph('💡 运动后注意拉伸，避免肌肉损伤', style='List Bullet')
    
    doc.add_page_break()
    
    # 第四部分：健康建议
    doc.add_heading('四、综合健康建议', level=1)
    
    doc.add_heading('1. 睡眠改善', level=2)
    doc.add_paragraph('• 保持规律作息，建议22:30前入睡', style='List Number')
    doc.add_paragraph('• 睡前可进行冥想或深呼吸练习', style='List Number')
    doc.add_paragraph('• 避免睡前饮用咖啡或浓茶', style='List Number')
    doc.add_paragraph('• 保持卧室安静、黑暗、凉爽', style='List Number')
    
    doc.add_heading('2. 运动计划', level=2)
    doc.add_paragraph('• 每周至少3次有氧运动，每次30-45分钟', style='List Number')
    doc.add_paragraph('• 每周2次力量训练，重点锻炼核心肌群', style='List Number')
    doc.add_paragraph('• 每日步行目标: 10,000步', style='List Number')
    doc.add_paragraph('• 运动前后做好热身和拉伸', style='List Number')
    
    doc.add_heading('3. 饮食建议', level=2)
    doc.add_paragraph('• 控制碳水摄入，增加优质蛋白质', style='List Number')
    doc.add_paragraph('• 多吃蔬菜水果，保证膳食纤维摄入', style='List Number')
    doc.add_paragraph('• 每日饮水至少2000ml', style='List Number')
    doc.add_paragraph('• 避免高糖、高油、高盐食物', style='List Number')
    
    doc.add_heading('4. 压力管理', level=2)
    doc.add_paragraph('• 适当放松，避免过度劳累', style='List Number')
    doc.add_paragraph('• 培养兴趣爱好，丰富业余生活', style='List Number')
    doc.add_paragraph('• 保持社交活动，与朋友家人多交流', style='List Number')
    doc.add_paragraph('• 学会时间管理，提高工作效率', style='List Number')
    
    doc.add_page_break()
    
    # 第五部分：下周目标
    doc.add_heading('五、下周健康目标', level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    headers = ['目标项目', '目标值', '当前状态']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    goals = [
        ('睡眠时长', '≥ 7.5小时', '待完成'),
        ('运动时长', '≥ 40分钟/天', '待完成'),
        ('步数', '≥ 9,000步/天', '待完成'),
        ('力量训练', '≥ 2次/周', '待完成')
    ]
    
    for i, goal in enumerate(goals, 1):
        for j, cell_data in enumerate(goal):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # 总结
    doc.add_heading('总结', level=1)
    p = doc.add_paragraph()
    p.add_run('本周整体健康表现良好，各项指标均有提升。睡眠质量稳定，运动量达标，步数有所增加。')
    p.add_run('建议下周继续保持良好习惯，重点关注力量训练和深度睡眠质量的提升。')
    
    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('报告生成时间: ').bold = True
    p.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存
    doc.save('str(PROJECT_ROOT)/skills/data-tracker/output/健康数据周报_完整版.docx')
    print("✅ 健康数据周报（完整版）已生成")

def create_diet_plan_with_nutrition():
    """创建带营养分析的减脂食谱"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('减脂食谱方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('专业营养师定制 · 科学减脂')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 用户信息
    p = doc.add_paragraph()
    p.add_run('用户: ').bold = True
    p.add_run('张三\n')
    p.add_run('目标: ').bold = True
    p.add_run('减脂\n')
    p.add_run('每日热量: ').bold = True
    p.add_run('1500千卡\n')
    p.add_run('蛋白质目标: ').bold = True
    p.add_run('120g (32%)\n')
    p.add_run('碳水目标: ').bold = True
    p.add_run('150g (40%)\n')
    p.add_run('脂肪目标: ').bold = True
    p.add_run('40g (24%)')
    
    doc.add_page_break()
    
    # 早餐
    doc.add_heading('早餐 (400千卡)', level=1)
    doc.add_heading('时间: 7:00-8:00', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    headers = ['食物', '份量', '热量']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    breakfast = [
        ('全麦面包', '2片', '160千卡'),
        ('水煮蛋', '2个', '140千卡'),
        ('黄瓜', '1根', '20千卡'),
        ('番茄', '1个', '20千卡'),
        ('黑咖啡', '1杯', '5千卡')
    ]
    
    for i, food in enumerate(breakfast, 1):
        for j, cell_data in enumerate(food):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_heading('营养说明', level=2)
    doc.add_paragraph('• 全麦面包提供优质碳水化合物，富含膳食纤维', style='List Bullet')
    doc.add_paragraph('• 水煮蛋是优质蛋白质来源，含有人体必需氨基酸', style='List Bullet')
    doc.add_paragraph('• 黄瓜和番茄富含维生素和矿物质，热量低', style='List Bullet')
    doc.add_paragraph('• 黑咖啡可提高新陈代谢，促进脂肪燃烧', style='List Bullet')
    
    doc.add_page_break()
    
    # 午餐
    doc.add_heading('午餐 (600千卡)', level=1)
    doc.add_heading('时间: 12:00-13:00', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    lunch = [
        ('糙米饭', '1碗', '180千卡'),
        ('鸡胸肉', '150g', '165千卡'),
        ('清蒸鱼', '100g', '120千卡'),
        ('西兰花', '200g', '60千卡'),
        ('紫菜蛋花汤', '1碗', '40千卡')
    ]
    
    for i, food in enumerate(lunch, 1):
        for j, cell_data in enumerate(food):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_heading('营养说明', level=2)
    doc.add_paragraph('• 糙米饭是低GI主食，可稳定血糖', style='List Bullet')
    doc.add_paragraph('• 鸡胸肉和鱼肉提供优质蛋白质，脂肪含量低', style='List Bullet')
    doc.add_paragraph('• 西兰花富含维生素C和膳食纤维', style='List Bullet')
    doc.add_paragraph('• 紫菜富含碘和矿物质，有助于代谢', style='List Bullet')
    
    doc.add_page_break()
    
    # 晚餐
    doc.add_heading('晚餐 (500千卡)', level=1)
    doc.add_heading('时间: 18:00-19:00', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    dinner = [
        ('红薯', '1个', '120千卡'),
        ('虾仁', '100g', '95千卡'),
        ('豆腐', '100g', '80千卡'),
        ('菠菜', '200g', '50千卡'),
        ('苹果', '1个', '95千卡')
    ]
    
    for i, food in enumerate(dinner, 1):
        for j, cell_data in enumerate(food):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_heading('营养说明', level=2)
    doc.add_paragraph('• 红薯是优质碳水，富含β-胡萝卜素', style='List Bullet')
    doc.add_paragraph('• 虾仁和豆腐提供优质蛋白质', style='List Bullet')
    doc.add_paragraph('• 菠菜富含铁和叶酸，有助于血液循环', style='List Bullet')
    doc.add_paragraph('• 苹果富含果胶，有助于肠道健康', style='List Bullet')
    
    doc.add_page_break()
    
    # 营养分析图表
    doc.add_heading('营养分析', level=1)
    
    # 生成营养饼图
    labels = ['蛋白质', '碳水化合物', '脂肪']
    sizes = [32, 40, 24]
    colors = ['#FF6B6B', '#4ECDC4', '#FFE66D']
    explode = (0.05, 0.05, 0.05)
    
    fig, ax = plt.subplots(figsize=(8, 8))
    ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%',
           shadow=True, startangle=90)
    ax.set_title('每日营养素占比', fontsize=16, fontweight='bold', pad=20)
    
    plt.tight_layout()
    nutrition_chart = 'str(PROJECT_ROOT)/skills/fitness-coach/output/nutrition_chart.png'
    plt.savefig(nutrition_chart, dpi=150, bbox_inches='tight')
    plt.close()
    
    add_chart_image(doc, nutrition_chart)
    
    # 营养素详细表格
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    headers = ['营养素', '含量', '占比', '评价']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    nutrients = [
        ('蛋白质', '120g', '32%', '✅ 达标'),
        ('碳水化合物', '150g', '40%', '✅ 达标'),
        ('脂肪', '40g', '24%', '✅ 达标'),
        ('纤维素', '25g', '-', '✅ 充足')
    ]
    
    for i, nutrient in enumerate(nutrients, 1):
        for j, cell_data in enumerate(nutrient):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_page_break()
    
    # 购物清单
    doc.add_heading('一周购物清单', level=1)
    
    doc.add_heading('蛋白质类', level=2)
    doc.add_paragraph('□ 鸡胸肉 500g', style='List Bullet')
    doc.add_paragraph('□ 鸡蛋 1盒（10个）', style='List Bullet')
    doc.add_paragraph('□ 虾仁 300g', style='List Bullet')
    doc.add_paragraph('□ 豆腐 2块', style='List Bullet')
    doc.add_paragraph('□ 鱼 300g（鲈鱼或鳕鱼）', style='List Bullet')
    
    doc.add_heading('蔬菜类', level=2)
    doc.add_paragraph('□ 西兰花 1颗', style='List Bullet')
    doc.add_paragraph('□ 菠菜 1把', style='List Bullet')
    doc.add_paragraph('□ 芹菜 2根', style='List Bullet')
    doc.add_paragraph('□ 黄瓜 3根', style='List Bullet')
    doc.add_paragraph('□ 番茄 5个', style='List Bullet')
    doc.add_paragraph('□ 胡萝卜 2根', style='List Bullet')
    
    doc.add_heading('主食类', level=2)
    doc.add_paragraph('□ 全麦面包 1袋', style='List Bullet')
    doc.add_paragraph('□ 糙米 1袋（500g）', style='List Bullet')
    doc.add_paragraph('□ 红薯 3个', style='List Bullet')
    
    doc.add_heading('水果类', level=2)
    doc.add_paragraph('□ 苹果 5个', style='List Bullet')
    doc.add_paragraph('□ 橙子 5个', style='List Bullet')
    
    doc.add_heading('其他', level=2)
    doc.add_paragraph('□ 紫菜 1包', style='List Bullet')
    doc.add_paragraph('□ 黑咖啡 1盒', style='List Bullet')
    
    doc.add_page_break()
    
    # 注意事项
    doc.add_heading('重要注意事项', level=1)
    
    doc.add_heading('烹饪方式', level=2)
    doc.add_paragraph('1. 优先选择清蒸、水煮、凉拌等低油烹饪方式', style='List Number')
    doc.add_paragraph('2. 炒菜时使用橄榄油或椰子油，控制用油量', style='List Number')
    doc.add_paragraph('3. 避免油炸、红烧等高油高盐烹饪方式', style='List Number')
    
    doc.add_heading('饮水建议', level=2)
    doc.add_paragraph('• 每日饮水量: 至少2000ml', style='List Bullet')
    doc.add_paragraph('• 最佳饮水时间: 起床后、餐前30分钟、运动后', style='List Bullet')
    doc.add_paragraph('• 避免饮用含糖饮料和碳酸饮料', style='List Bullet')
    
    doc.add_heading('运动配合', level=2)
    doc.add_paragraph('• 每周至少3-5次有氧运动，每次30-45分钟', style='List Bullet')
    doc.add_paragraph('• 推荐运动: 快走、慢跑、游泳、骑行', style='List Bullet')
    doc.add_paragraph('• 运动时间: 早餐前或晚餐后1小时', style='List Bullet')
    
    doc.add_heading('作息建议', level=2)
    doc.add_paragraph('• 规律作息，避免熬夜', style='List Bullet')
    doc.add_paragraph('• 保证每日7-8小时睡眠', style='List Bullet')
    doc.add_paragraph('• 睡前3小时避免进食', style='List Bullet')
    
    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('方案制定: ').bold = True
    p.add_run('AI营养师\n')
    p.add_run('生成时间: ').bold = True
    p.add_run(datetime.now().strftime('%Y年%m月%d日'))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存
    doc.save('str(PROJECT_ROOT)/skills/fitness-coach/output/减脂食谱_完整版.docx')
    print("✅ 减脂食谱（完整版）已生成")

if __name__ == '__main__':
    create_health_report_with_charts()
    create_diet_plan_with_nutrition()
    print("\n所有完整版文档已生成！")
