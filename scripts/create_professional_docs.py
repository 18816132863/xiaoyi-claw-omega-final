from pathlib import Path
#!/usr/bin/env python3

PROJECT_ROOT = Path(__file__).resolve().parents[1]
"""
生成专业格式的文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

# 注册中文字体
font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
if os.path.exists(font_path):
    pdfmetrics.registerFont(TTFont('DejaVu', font_path))

def create_weekly_report():
    """创建周报Word文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('周报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('报告人: ').bold = True
    p.add_run('张三\n')
    p.add_run('部门: ').bold = True
    p.add_run('技术部\n')
    p.add_run('时间: ').bold = True
    p.add_run('2026年4月15日 - 2026年4月19日')
    
    # 本周工作
    doc.add_heading('本周工作完成情况', level=1)
    
    doc.add_heading('1. 技能升级系统开发', level=2)
    doc.add_paragraph('• 完成了11个核心技能的商业级别升级', style='List Bullet')
    doc.add_paragraph('• 实现了统一的升级框架和机制', style='List Bullet')
    doc.add_paragraph('• 创建了批量升级脚本', style='List Bullet')
    doc.add_paragraph('• 所有技能通过测试验证', style='List Bullet')
    
    doc.add_heading('2. 架构优化', level=2)
    doc.add_paragraph('• 建立了六层架构体系', style='List Bullet')
    doc.add_paragraph('• 完善了层间依赖规则', style='List Bullet')
    doc.add_paragraph('• 创建了质量门禁系统', style='List Bullet')
    doc.add_paragraph('• 实现了审计日志记录', style='List Bullet')
    
    # 下周计划
    doc.add_heading('下周工作计划', level=1)
    
    doc.add_heading('1. 功能增强', level=2)
    doc.add_paragraph('• 集成主流API（绘图、音乐、视频）', style='List Bullet')
    doc.add_paragraph('• 添加更多模板和场景', style='List Bullet')
    doc.add_paragraph('• 优化用户体验', style='List Bullet')
    
    # 问题与建议
    doc.add_heading('问题与建议', level=1)
    
    doc.add_heading('存在问题', level=2)
    doc.add_paragraph('• 部分技能需要进一步优化', style='List Bullet')
    doc.add_paragraph('• API集成需要完善', style='List Bullet')
    
    doc.add_heading('改进建议', level=2)
    doc.add_paragraph('• 加强测试覆盖', style='List Bullet')
    doc.add_paragraph('• 完善文档说明', style='List Bullet')
    doc.add_paragraph('• 优化用户引导', style='List Bullet')
    
    # 保存
    doc.save('str(PROJECT_ROOT)/skills/doc-autofill/output/周报_20260419.docx')
    print("✅ 周报Word文档已生成")

def create_novel_chapter():
    """创建小说章节Word文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('第1章 重生归来', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期
    date = doc.add_paragraph('2026年4月18日 晴')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 正文
    paragraphs = [
        '林峰睁开眼睛，发现自己躺在熟悉的出租屋里。墙上贴着泛黄的海报，桌上堆满了泡面盒和可乐瓶。',
        '"我...重生了？"',
        '林峰难以置信地看着自己的双手，年轻、有力，没有那道车祸留下的伤疤。他猛地坐起身，看向床头的日历——2016年4月18日。',
        '十年前。',
        '那个改变他命运的日子。',
        '林峰深吸一口气，脑海中浮现出前世的记忆：创业失败、负债累累、众叛亲离，最后在一场车祸中结束了自己短暂的一生。',
        '"这一次，我绝不会重蹈覆辙。"',
        '林峰握紧拳头，眼中闪过一丝坚定。他拥有未来十年的记忆，知道哪些行业会崛起，哪些股票会暴涨，哪些机会稍纵即逝。',
        '这一次，他要站在巅峰，俯瞰众生。',
        '手机突然响起，是母亲打来的电话。',
        '"小峰啊，你爸的手术费还差十万，你那边能想办法吗？"',
        '林峰心中一痛。前世，正是因为拿不出这笔钱，父亲的病情一拖再拖，最终...',
        '"妈，别担心，钱的事我来想办法。"',
        '挂断电话，林峰看向窗外。阳光正好，微风不燥。',
        '新的生活，开始了。'
    ]
    
    for para in paragraphs:
        doc.add_paragraph(para)
    
    # 统计信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('字数: ').bold = True
    p.add_run('312字\n')
    p.add_run('爽点: ').bold = True
    p.add_run('重生、金手指、逆袭\n')
    p.add_run('下章预告: ').bold = True
    p.add_run('林峰如何快速赚到第一桶金？')
    
    # 保存
    doc.save('str(PROJECT_ROOT)/skills/novel-generator/output/第001章_重生归来.docx')
    print("✅ 小说章节Word文档已生成")

def create_diet_plan():
    """创建减脂食谱Word文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('减脂食谱 - 一日三餐', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('目标: ').bold = True
    p.add_run('减脂\n')
    p.add_run('每日热量: ').bold = True
    p.add_run('1500千卡\n')
    p.add_run('餐数: ').bold = True
    p.add_run('3餐')
    
    # 早餐
    doc.add_heading('早餐 (400千卡)', level=1)
    
    doc.add_heading('主食', level=2)
    doc.add_paragraph('全麦面包 2片 (160千卡)', style='List Bullet')
    
    doc.add_heading('蛋白质', level=2)
    doc.add_paragraph('水煮蛋 2个 (140千卡)', style='List Bullet')
    
    doc.add_heading('蔬菜', level=2)
    doc.add_paragraph('黄瓜 1根 (20千卡)', style='List Bullet')
    doc.add_paragraph('番茄 1个 (20千卡)', style='List Bullet')
    
    # 午餐
    doc.add_heading('午餐 (600千卡)', level=1)
    
    doc.add_heading('主食', level=2)
    doc.add_paragraph('糙米饭 1碗 (180千卡)', style='List Bullet')
    
    doc.add_heading('蛋白质', level=2)
    doc.add_paragraph('鸡胸肉 150g (165千卡)', style='List Bullet')
    doc.add_paragraph('清蒸鱼 100g (120千卡)', style='List Bullet')
    
    # 晚餐
    doc.add_heading('晚餐 (500千卡)', level=1)
    
    doc.add_heading('主食', level=2)
    doc.add_paragraph('红薯 1个 (120千卡)', style='List Bullet')
    
    doc.add_heading('蛋白质', level=2)
    doc.add_paragraph('虾仁 100g (95千卡)', style='List Bullet')
    doc.add_paragraph('豆腐 100g (80千卡)', style='List Bullet')
    
    # 营养分析
    doc.add_heading('营养分析', level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '营养素'
    hdr_cells[1].text = '含量'
    hdr_cells[2].text = '占比'
    
    # 数据
    data = [
        ('蛋白质', '120g', '32%'),
        ('碳水化合物', '150g', '40%'),
        ('脂肪', '40g', '24%'),
        ('纤维素', '25g', '-')
    ]
    
    for i, (nutrient, amount, percent) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = nutrient
        row_cells[1].text = amount
        row_cells[2].text = percent
    
    # 购物清单
    doc.add_heading('购物清单', level=1)
    
    doc.add_heading('蛋白质类', level=2)
    doc.add_paragraph('□ 鸡胸肉 500g', style='List Bullet')
    doc.add_paragraph('□ 鸡蛋 1盒', style='List Bullet')
    doc.add_paragraph('□ 虾仁 300g', style='List Bullet')
    
    # 注意事项
    doc.add_heading('注意事项', level=1)
    doc.add_paragraph('1. 烹饪方式: 清蒸、水煮、少油少盐', style='List Number')
    doc.add_paragraph('2. 饮水: 每日至少2000ml', style='List Number')
    doc.add_paragraph('3. 运动: 配合有氧运动，每周3-5次', style='List Number')
    doc.add_paragraph('4. 作息: 规律作息，避免熬夜', style='List Number')
    
    # 保存
    doc.save('str(PROJECT_ROOT)/skills/fitness-coach/output/减脂食谱_20260418.docx')
    print("✅ 减脂食谱Word文档已生成")

if __name__ == '__main__':
    create_weekly_report()
    create_novel_chapter()
    create_diet_plan()
    print("\n所有Word文档已生成！")
