#!/usr/bin/env python3
"""
Doc Autofill - 文档自动填写系统 V1.0.0

功能：
1. 解析模板占位符
2. 从长期记忆读取数据
3. 自动填写文档
4. 批量生成
"""

import json
import sys
import re
from pathlib import Path
from datetime import datetime, date
from typing import Dict, List, Any, Optional


def get_project_root() -> Path:
    """获取项目根目录"""
    current = Path(__file__).resolve().parent.parent.parent
    if (current / 'core' / 'ARCHITECTURE.md').exists():
        return current
    return Path(__file__).resolve().parent.parent.parent


class DocAutofill:
    """文档自动填写器"""
    
    def __init__(self, root: Path = None):
        self.root = root or get_project_root()
        self.template_dir = self.root / "reports" / "templates"
        self.output_dir = self.root / "reports" / "output"
        self.memory_dir = self.root / "reports" / "live_learning"
        
        self.template_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 占位符模式
        self.placeholder_pattern = re.compile(r'\{\{(\w+)\}\}')
        
        # 字段映射
        self.field_mapping = {
            "姓名": "name", "名字": "name",
            "电话": "phone", "手机": "phone", "联系电话": "phone",
            "邮箱": "email", "电子邮件": "email",
            "身份证": "id_card", "身份证号": "id_card",
            "公司": "company", "单位": "company",
            "地址": "address", "住址": "address",
            "日期": "date", "时间": "date",
            "银行": "bank_account", "银行卡": "bank_account"
        }
    
    def load_personal_info(self) -> Dict:
        """加载个人信息"""
        info_file = self.memory_dir / "user_preferences.json"
        if info_file.exists():
            data = json.loads(info_file.read_text(encoding='utf-8'))
            return data.get("personal_info", {})
        return {}
    
    def save_personal_info(self, info: Dict):
        """保存个人信息"""
        info_file = self.memory_dir / "user_preferences.json"
        
        existing = {}
        if info_file.exists():
            existing = json.loads(info_file.read_text(encoding='utf-8'))
        
        existing["personal_info"] = info
        info_file.write_text(json.dumps(existing, ensure_ascii=False, indent=2), encoding='utf-8')
    
    def parse_template(self, template_path: Path) -> Dict:
        """解析模板"""
        result = {
            "path": str(template_path),
            "placeholders": [],
            "fields": []
        }
        
        if template_path.suffix == '.docx':
            result["placeholders"] = self._parse_docx(template_path)
        elif template_path.suffix == '.xlsx':
            result["placeholders"] = self._parse_xlsx(template_path)
        elif template_path.suffix in ['.md', '.txt']:
            result["placeholders"] = self._parse_text(template_path)
        
        # 转换为字段名
        for ph in result["placeholders"]:
            field = self.field_mapping.get(ph, ph.lower())
            if field not in result["fields"]:
                result["fields"].append(field)
        
        return result
    
    def _parse_docx(self, path: Path) -> List[str]:
        """解析 Word 文档"""
        placeholders = []
        try:
            from docx import Document
            doc = Document(str(path))
            
            # 检查段落
            for para in doc.paragraphs:
                matches = self.placeholder_pattern.findall(para.text)
                placeholders.extend(matches)
            
            # 检查表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = self.placeholder_pattern.findall(cell.text)
                        placeholders.extend(matches)
        except ImportError:
            pass
        except Exception as e:
            print(f"解析 Word 失败: {e}")
        
        return list(set(placeholders))
    
    def _parse_xlsx(self, path: Path) -> List[str]:
        """解析 Excel 文档"""
        placeholders = []
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path))
            
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            matches = self.placeholder_pattern.findall(str(cell.value))
                            placeholders.extend(matches)
        except ImportError:
            pass
        except Exception as e:
            print(f"解析 Excel 失败: {e}")
        
        return list(set(placeholders))
    
    def _parse_text(self, path: Path) -> List[str]:
        """解析文本文件"""
        content = path.read_text(encoding='utf-8')
        matches = self.placeholder_pattern.findall(content)
        return list(set(matches))
    
    def fill_template(
        self,
        template_path: Path,
        data: Dict,
        output_name: str = None
    ) -> Dict:
        """填写模板"""
        result = {
            "success": False,
            "output_path": None,
            "filled_fields": [],
            "missing_fields": []
        }
        
        # 检查缺失字段
        template_info = self.parse_template(template_path)
        for field in template_info["fields"]:
            if field not in data:
                result["missing_fields"].append(field)
        
        if template_path.suffix == '.docx':
            output_path = self._fill_docx(template_path, data, output_name)
        elif template_path.suffix == '.xlsx':
            output_path = self._fill_xlsx(template_path, data, output_name)
        elif template_path.suffix in ['.md', '.txt']:
            output_path = self._fill_text(template_path, data, output_name)
        else:
            return result
        
        if output_path:
            result["success"] = True
            result["output_path"] = str(output_path)
            result["filled_fields"] = [k for k in data.keys() if k in template_info["fields"]]
        
        return result
    
    def _fill_docx(self, template_path: Path, data: Dict, output_name: str = None) -> Optional[Path]:
        """填写 Word 文档"""
        try:
            from docx import Document
            doc = Document(str(template_path))
            
            # 替换段落
            for para in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(value))
            
            # 替换表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            # 保存
            if not output_name:
                output_name = f"{template_path.stem}_已填写.docx"
            output_path = self.output_dir / output_name
            doc.save(str(output_path))
            
            return output_path
        except Exception as e:
            print(f"填写 Word 失败: {e}")
            return None
    
    def _fill_xlsx(self, template_path: Path, data: Dict, output_name: str = None) -> Optional[Path]:
        """填写 Excel 文档"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(template_path))
            
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            for key, value in data.items():
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in str(cell.value):
                                    cell.value = str(cell.value).replace(placeholder, str(value))
            
            # 保存
            if not output_name:
                output_name = f"{template_path.stem}_已填写.xlsx"
            output_path = self.output_dir / output_name
            wb.save(str(output_path))
            
            return output_path
        except Exception as e:
            print(f"填写 Excel 失败: {e}")
            return None
    
    def _fill_text(self, template_path: Path, data: Dict, output_name: str = None) -> Optional[Path]:
        """填写文本文件"""
        try:
            content = template_path.read_text(encoding='utf-8')
            
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                content = content.replace(placeholder, str(value))
            
            # 保存
            if not output_name:
                output_name = f"{template_path.stem}_已填写{template_path.suffix}"
            output_path = self.output_dir / output_name
            output_path.write_text(content, encoding='utf-8')
            
            return output_path
        except Exception as e:
            print(f"填写文本失败: {e}")
            return None


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Doc Autofill V1.0.0")
    parser.add_argument("action", choices=["parse", "fill", "info"], help="操作类型")
    parser.add_argument("--template", help="模板文件路径")
    parser.add_argument("--data", help="数据JSON")
    args = parser.parse_args()
    
    root = get_project_root()
    autofill = DocAutofill(root)
    
    if args.action == "parse":
        if not args.template:
            print("请提供模板路径: --template path/to/template.docx")
            return 1
        
        result = autofill.parse_template(Path(args.template))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        
    elif args.action == "fill":
        if not args.template:
            print("请提供模板路径: --template path/to/template.docx")
            return 1
        
        # 合并个人信息和用户数据
        data = autofill.load_personal_info()
        if args.data:
            user_data = json.loads(args.data)
            data.update(user_data)
        
        result = autofill.fill_template(Path(args.template), data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        
    elif args.action == "info":
        info = autofill.load_personal_info()
        print(json.dumps(info, ensure_ascii=False, indent=2))
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
