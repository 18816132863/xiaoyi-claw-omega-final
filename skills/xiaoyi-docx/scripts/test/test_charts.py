import sys
import os
import tempfile
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from matplotlib_chart_utils import (
    add_column_chart,
    add_bar_chart,
    add_line_chart,
    add_pie_chart,
    add_area_chart,
    add_scatter_chart,
    add_radar_chart,
    add_histogram,
)


def _assert_docx_has_media(doc_path):
    import zipfile
    with zipfile.ZipFile(doc_path, "r") as zf:
        media_files = [n for n in zf.namelist() if "media" in n]
    assert len(media_files) >= 1, f"Expected at least one media file, got {media_files}"


def test_add_column_chart():
    doc = Document()
    add_column_chart(doc, categories=["A", "B", "C"], series=[{"name": "S1", "values": [1, 2, 3]}], title="Column")
    path = os.path.join(tempfile.gettempdir(), "test_column.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_bar_chart():
    doc = Document()
    add_bar_chart(doc, categories=["X", "Y", "Z"], series=[{"name": "S1", "values": [3, 2, 1]}], title="Bar")
    path = os.path.join(tempfile.gettempdir(), "test_bar.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_line_chart():
    doc = Document()
    add_line_chart(doc, categories=["1", "2", "3"], series=[{"name": "S1", "values": [1, 3, 2]}], title="Line")
    path = os.path.join(tempfile.gettempdir(), "test_line.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_pie_chart():
    doc = Document()
    add_pie_chart(doc, categories=["A", "B", "C"], series=[{"name": "S1", "values": [30, 40, 30]}], title="Pie")
    path = os.path.join(tempfile.gettempdir(), "test_pie.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_area_chart():
    doc = Document()
    add_area_chart(doc, categories=["Q1", "Q2"], series=[{"name": "S1", "values": [4, 7]}], title="Area")
    path = os.path.join(tempfile.gettempdir(), "test_area.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_scatter_chart():
    doc = Document()
    add_scatter_chart(
        doc,
        series=[
            {"name": "S1", "values": [(1, 3), (2, 5), (3, 2)]},
            {"name": "S2", "values": [(1, 4), (2, 2), (3, 6)]},
        ],
        title="Scatter",
    )
    path = os.path.join(tempfile.gettempdir(), "test_scatter.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_radar_chart():
    doc = Document()
    add_radar_chart(doc, categories=["A", "B", "C"], series=[{"name": "S1", "values": [3, 4, 5]}], title="Radar")
    path = os.path.join(tempfile.gettempdir(), "test_radar.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_add_histogram():
    doc = Document()
    add_histogram(doc, data=[1, 2, 2, 3, 3, 3, 4, 4, 5], bins=5, title="Histogram")
    path = os.path.join(tempfile.gettempdir(), "test_histogram.docx")
    doc.save(path)
    _assert_docx_has_media(path)


def test_series_validation():
    doc = Document()
    try:
        add_column_chart(doc, categories=["A", "B"], series=[{"name": "S1", "values": [1]}])
        assert False, "Expected ValueError"
    except ValueError as e:
        assert "does not match categories length" in str(e)
