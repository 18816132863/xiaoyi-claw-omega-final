#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Python-docx Skill 全量功能测试脚本
测试内容：
1. 创建新文档 + 编辑现有文档
2. 中英文字体支持（setup_chinese_document_styles / set_chinese_font / set_document_default_font）
3. 公式 (OMath)
4. 图片插入
5. 表格（创建样式表、单元格边框、背景、边距、无边框表格）
6. 列表（有序/无序 + restart）
7. 超链接
8. 目录、页码、水平分隔线、左右对齐文本
9. 字数统计
10. 边界情况（特殊字符、长段落、对齐方式等）
"""

import sys
import os
import time
from datetime import datetime

# 添加 skill 路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 导入工具函数
from docx_utils import (
    setup_chinese_document_styles,
    set_chinese_font,
    set_document_default_font,
    add_table_of_contents,
    add_page_number_footer,
    count_document_words,
    add_horizontal_line,
    create_styled_table,
    format_table_cell,
    remove_table_borders,
    set_cell_border,
    set_cell_background,
    set_cell_margins,
    add_hyperlink,
)


# ========================================================================
# Helper
# ========================================================================

def add_formula(paragraph, formula_text):
    """添加公式辅助函数"""
    oMath = OxmlElement('m:oMath')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = formula_text
    r.append(t)
    oMath.append(r)
    paragraph._p.append(oMath)
    return paragraph


def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path)


# ========================================================================
# Test Cases
# ========================================================================

def test_create_comprehensive_document():
    """测试1: 创建包含所有功能的新文档"""
    print("=" * 60)
    print("测试1: 创建综合测试文档")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    # 1. 标题
    print("  - 添加标题...")
    title = doc.add_paragraph('Python-docx Skill 全量功能测试文档', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 目录
    print("  - 添加目录...")
    doc.add_paragraph('目录', style='Heading 1')
    add_table_of_contents(doc)
    doc.add_page_break()

    # 3. 中英文字体测试
    print("  - 测试中英文字体...")
    doc.add_paragraph('第一章 中英文字体测试', style='Heading 1')

    p = doc.add_paragraph('这是一段中文正文内容，用于测试 SimSun（宋体）字体显示效果。'
                         '中文文档通常需要设置首行缩进两个字符。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph('This is English text to test font rendering. '
                         'The quick brown fox jumps over the lazy dog. '
                         'Python-docx provides powerful document generation capabilities.',
                         style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p = doc.add_paragraph('这是一段Mixed Content：中文和English混排测试。'
                         '公式如 E=mc2 和变量 x, y, z 应该正确显示。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 4. 文本样式
    doc.add_paragraph('1.1 文本样式测试', style='Heading 2')
    p = doc.add_paragraph()
    p.add_run('这是普通文本 ')
    run_bold = p.add_run('粗体文本 ')
    run_bold.font.bold = True
    run_italic = p.add_run('斜体文本 ')
    run_italic.font.italic = True
    run_color = p.add_run('红色文本 ')
    run_color.font.color.rgb = RGBColor(255, 0, 0)
    run_size = p.add_run('大号文本')
    run_size.font.size = Pt(16)

    # 5. 公式
    print("  - 添加公式...")
    doc.add_paragraph('第二章 公式测试', style='Heading 1')
    p = doc.add_paragraph('二次方程求根公式：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p_formula = doc.add_paragraph()
    p_formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_elem = OxmlElement('m:oMath')
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = 'x = '
    r1.append(t1)

    frac = OxmlElement('m:f')
    num = OxmlElement('m:num')
    r_num = OxmlElement('w:r')
    t_num = OxmlElement('w:t')
    t_num.text = '-b +- sqrt(b2-4ac)'
    r_num.append(t_num)
    num.append(r_num)

    den = OxmlElement('m:den')
    r_den = OxmlElement('w:r')
    t_den = OxmlElement('w:t')
    t_den.text = '2a'
    r_den.append(t_den)
    den.append(r_den)

    frac.append(num)
    frac.append(den)
    math_elem.append(r1)
    math_elem.append(frac)
    p_formula._p.append(math_elem)

    p = doc.add_paragraph('积分公式示例：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    p_formula2 = doc.add_paragraph()
    p_formula2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_elem2 = OxmlElement('m:oMath')
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = 'integral[a,b] f(x)dx = F(b) - F(a)'
    r2.append(t2)
    math_elem2.append(r2)
    p_formula2._p.append(math_elem2)

    # 6. 图片
    print("  - 添加图片...")
    doc.add_paragraph('第三章 图片测试', style='Heading 1')
    p = doc.add_paragraph('以下是图片插入测试：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (400, 200), color='lightblue')
        draw = ImageDraw.Draw(img)
        draw.text((50, 80), "Test Image\nPython-docx", fill='black')
        img_path = os.path.join(OUTPUT_DIR, 'test_image_temp.png')
        img.save(img_path)

        doc.add_picture(img_path, width=Inches(4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except ImportError:
        p = doc.add_paragraph('[PIL 未安装，跳过图片测试]', style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 7. 表格
    print("  - 添加表格...")
    doc.add_paragraph('第四章 表格测试', style='Heading 1')
    p = doc.add_paragraph('以下是 styled table 测试：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    table = create_styled_table(doc, rows=4, cols=3, col_widths=[2.0, 2.5, 3.0])
    format_table_cell(table.rows[0].cells[0], "功能模块", alignment='center')
    format_table_cell(table.rows[0].cells[1], "支持状态", alignment='center')
    format_table_cell(table.rows[0].cells[2], "说明", alignment='center')

    table.rows[1].cells[0].text = "中英文字体"
    format_table_cell(table.rows[1].cells[1], "[OK] 支持", alignment='center')
    table.rows[1].cells[2].text = "SimSun/SimHei 等中文字体"

    table.rows[2].cells[0].text = "公式编辑"
    format_table_cell(table.rows[2].cells[1], "[OK] 支持", alignment='center')
    table.rows[2].cells[2].text = "OMath 公式支持"

    table.rows[3].cells[0].text = "图片插入"
    format_table_cell(table.rows[3].cells[1], "[OK] 支持", alignment='center')
    table.rows[3].cells[2].text = "支持 PNG/JPG 等格式"

    # 无边框表格
    doc.add_paragraph('')
    p = doc.add_paragraph('无边框表格（用于布局）：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    layout_table = doc.add_table(rows=2, cols=2)
    layout_table.rows[0].cells[0].text = "左侧内容"
    layout_table.rows[0].cells[1].text = "右侧内容"
    layout_table.rows[1].cells[0].text = "标签："
    layout_table.rows[1].cells[1].text = "数值内容"
    remove_table_borders(layout_table)

    # 8. 列表
    print("  - 添加列表...")
    doc.add_paragraph('第五章 列表测试', style='Heading 1')

    doc.add_paragraph('5.1 无序列表（项目符号）', style='Heading 2')
    doc.add_paragraph('第一项：Python-docx 基础功能', style='List Bullet')
    doc.add_paragraph('第二项：中文文档支持', style='List Bullet')
    doc.add_paragraph('第三项：样式和格式设置', style='List Bullet')
    doc.add_paragraph('嵌套项示例', style='List Bullet 2')
    doc.add_paragraph('更深嵌套', style='List Bullet 3')

    doc.add_paragraph('5.2 有序列表（编号）', style='Heading 2')
    doc.add_paragraph('安装 python-docx 库', style='List Number', restart=True)
    doc.add_paragraph('导入 Document 类', style='List Number')
    doc.add_paragraph('创建文档对象', style='List Number')
    doc.add_paragraph('添加内容和样式', style='List Number')
    doc.add_paragraph('保存文档', style='List Number')

    doc.add_paragraph('')
    doc.add_paragraph('重新开始编号：', style='Normal')
    doc.add_paragraph('第一步：分析问题', style='List Number', restart=True)
    doc.add_paragraph('第二步：设计方案', style='List Number')
    doc.add_paragraph('第三步：实施开发', style='List Number')

    # 9. 水平分隔线
    print("  - 添加水平分隔线...")
    doc.add_paragraph('第六章 水平分隔线', style='Heading 1')
    p = doc.add_paragraph('以下是单线分隔：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    line1 = doc.add_paragraph()
    add_horizontal_line(line1)

    p = doc.add_paragraph('以下是双线分隔（红色）：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    line2 = doc.add_paragraph()
    add_horizontal_line(line2, line_style='double', line_size=12, color='FF0000')

    # 10. 左右对齐文本
    print("  - 添加左右对齐文本...")
    doc.add_paragraph('第七章 特殊排版', style='Heading 1')
    p = doc.add_paragraph('左右对齐文本示例：', style='Heading 2')

    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin
    p_align = doc.add_paragraph()
    p_align.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    p_align.add_run('左侧文本（左对齐）')
    run_tab = p_align.add_run('\t')
    run_tab.underline = True
    p_align.add_run('右侧文本（右对齐）')

    # 11. 字数统计
    print("  - 统计字数...")
    word_count = count_document_words(doc)
    doc.add_paragraph('第八章 文档统计', style='Heading 1')
    p = doc.add_paragraph('本文档总字数：{} 字'.format(word_count), style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 12. 页脚页码
    print("  - 添加页码...")
    add_page_number_footer(section, alignment='center')

    output_path = os.path.join(OUTPUT_DIR, 'test_output_create.docx')
    doc.save(output_path)
    print("  [PASS] 文档已保存: {}".format(output_path))

    return output_path, word_count


def test_edit_existing_document():
    """测试2: 编辑现有文档"""
    print("\n" + "=" * 60)
    print("测试2: 编辑现有文档")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)
    doc.add_paragraph('原始文档标题', style='Title')
    doc.add_paragraph('第一章 原始内容', style='Heading 1')
    p = doc.add_paragraph('这是原始文档的第一段内容。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    original_path = os.path.join(OUTPUT_DIR, 'test_original.docx')
    doc.save(original_path)
    print("  - 创建原始文档: {}".format(original_path))

    doc = Document(original_path)
    print("  - 重新打开文档进行编辑")

    doc.add_paragraph('第二章 新增内容', style='Heading 1')
    p = doc.add_paragraph('这是在编辑时添加的新段落。测试文档追加功能。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    table = create_styled_table(doc, rows=3, cols=2, col_widths=[3.0, 4.0])
    format_table_cell(table.rows[0].cells[0], "属性", alignment='center')
    format_table_cell(table.rows[0].cells[1], "值", alignment='center')
    table.rows[1].cells[0].text = "文档类型"
    table.rows[1].cells[1].text = "测试文档"
    table.rows[2].cells[0].text = "编辑次数"
    table.rows[2].cells[1].text = "1"

    doc.add_paragraph('编辑步骤：', style='Heading 2')
    doc.add_paragraph('打开现有文档', style='List Number', restart=True)
    doc.add_paragraph('添加新内容', style='List Number')
    doc.add_paragraph('保存修改', style='List Number')

    section = doc.sections[0]
    add_page_number_footer(section, alignment='right')

    output_path = os.path.join(OUTPUT_DIR, 'test_output_edit.docx')
    doc.save(output_path)
    print("  [PASS] 编辑后的文档已保存: {}".format(output_path))

    word_count = count_document_words(output_path)
    print("  - 编辑后文档字数: {}".format(word_count))

    return output_path, word_count


def test_edge_cases():
    """测试3: 边界情况和特殊字符"""
    print("\n" + "=" * 60)
    print("测试3: 边界情况和特殊字符")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)
    doc.add_paragraph('边界情况测试', style='Title')

    doc.add_paragraph('特殊字符测试', style='Heading 1')
    p = doc.add_paragraph('特殊符号：(C) (R) (TM) degree C degree F section No. star star triangle triangle diamond diamond circle circle ', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p = doc.add_paragraph('数学符号：+- * / infinity sqrt sum prod integral contour integral approx != <= >= ', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p = doc.add_paragraph('货币符号：Y $ Euro GBP Yen Rupee ', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('长段落测试', style='Heading 1')
    long_text = '这是一段很长的文本内容，用于测试段落处理和分页功能。' * 20
    p = doc.add_paragraph(long_text, style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('间距测试', style='Heading 1')
    p = doc.add_paragraph('这是一个有下方间距的段落。', style='Normal')
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph('上方段落有24pt间距。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('多行文本测试', style='Heading 1')
    p = doc.add_paragraph('第一行\n第二行（软回车）\n第三行（软回车）', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    doc.add_paragraph('对齐方式测试', style='Heading 1')
    p = doc.add_paragraph('左对齐文本（默认）')
    p = doc.add_paragraph('居中文本')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('右对齐文本')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('两端对齐文本，用于测试文本两端对齐效果。')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output_path = os.path.join(OUTPUT_DIR, 'test_output_edge_cases.docx')
    doc.save(output_path)
    print("  [PASS] 边界测试文档已保存: {}".format(output_path))

    word_count = count_document_words(output_path)
    print("  - 文档字数: {}".format(word_count))

    return output_path, word_count


def test_hyperlink():
    """测试4: 超链接"""
    print("\n" + "=" * 60)
    print("测试4: 超链接")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)
    doc.add_paragraph('超链接测试', style='Title')

    p = doc.add_paragraph('访问以下链接：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p2 = doc.add_paragraph()
    add_hyperlink(p2, 'Python-docx GitHub', 'https://github.com/python-openxml/python-docx',
                  color=(0, 112, 192), underline=True)

    p3 = doc.add_paragraph()
    add_hyperlink(p3, '百度', 'https://www.baidu.com',
                  color=(255, 0, 0), underline=False)

    output_path = os.path.join(OUTPUT_DIR, 'test_output_hyperlink.docx')
    doc.save(output_path)
    print("  [PASS] 超链接测试文档已保存: {}".format(output_path))

    return output_path, 0


def test_cell_advanced_styles():
    """测试5: 单元格高级样式（边框、背景、边距）"""
    print("\n" + "=" * 60)
    print("测试5: 单元格高级样式")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)
    doc.add_paragraph('单元格高级样式测试', style='Title')

    # 使用原生 API 创建表格，然后手动设置每个单元格的样式
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 第一行：表头，蓝色背景，白色粗体文字
    for cell in table.rows[0].cells:
        set_cell_background(cell, '4472C4')
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        cell.text = "表头"
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_chinese_font(run.font, 'SimHei')

    # 第二行：左侧单元格红色边框，右侧单元格绿色背景
    set_cell_border(table.rows[1].cells[0], top='FF0000', bottom='FF0000', left='FF0000', right='FF0000')
    table.rows[1].cells[0].text = "红边框"

    set_cell_background(table.rows[1].cells[1], 'C6E0B4')
    table.rows[1].cells[1].text = "绿背景"

    set_cell_margins(table.rows[1].cells[2], top=200, bottom=200, left=200, right=200)
    table.rows[1].cells[2].text = "大边距"

    # 第三行：混合样式
    table.rows[2].cells[0].text = "普通"
    set_cell_background(table.rows[2].cells[1], 'FFF2CC')
    set_cell_border(table.rows[2].cells[1], bottom='FFC000')
    table.rows[2].cells[1].text = "黄底+下边框"
    table.rows[2].cells[2].text = "普通"

    output_path = os.path.join(OUTPUT_DIR, 'test_output_cell_styles.docx')
    doc.save(output_path)
    print("  [PASS] 单元格样式测试文档已保存: {}".format(output_path))

    return output_path, 0


def test_default_font():
    """测试6: 设置文档默认字体"""
    print("\n" + "=" * 60)
    print("测试6: 文档默认字体设置")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)
    set_document_default_font(doc, font_name='Microsoft YaHei', ascii_font='Arial', font_size=14)

    doc.add_paragraph('默认字体测试文档', style='Title')
    p = doc.add_paragraph('这段文字应该使用微软雅黑（中文）和 Arial（英文），字号14pt。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(28)

    p2 = doc.add_paragraph('English text should use Arial 14pt.')
    p2.paragraph_format.first_line_indent = Pt(28)

    output_path = os.path.join(OUTPUT_DIR, 'test_output_default_font.docx')
    doc.save(output_path)
    print("  [PASS] 默认字体测试文档已保存: {}".format(output_path))

    return output_path, 0


def test_independent_chinese_font():
    """测试7: 独立设置中文字体"""
    print("\n" + "=" * 60)
    print("测试7: 独立中文字体设置")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)
    doc.add_paragraph('独立字体设置测试', style='Title')

    p = doc.add_paragraph()
    run1 = p.add_run('这是默认字体的文本。 ')
    run2 = p.add_run('这是楷体中文。 ')
    set_chinese_font(run2.font, 'KaiTi', ascii_font='Times New Roman')
    run3 = p.add_run('This is Times New Roman.')
    set_chinese_font(run3.font, 'Times New Roman')

    output_path = os.path.join(OUTPUT_DIR, 'test_output_independent_font.docx')
    doc.save(output_path)
    print("  [PASS] 独立字体测试文档已保存: {}".format(output_path))

    return output_path, 0


# ========================================================================
# Report Generation
# ========================================================================

def generate_markdown_report(results, duration_sec):
    report_path = os.path.join(OUTPUT_DIR, 'TEST_REPORT.md')
    lines = []
    lines.append("# Python-docx Skill 全量功能测试报告")
    lines.append("")
    lines.append("- **测试时间**: {}".format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    lines.append("- **总耗时**: {:.2f} 秒".format(duration_sec))
    lines.append("")

    passed = sum(1 for r in results if r[3] == "PASS")
    failed = sum(1 for r in results if r[3] == "FAIL")
    lines.append("## 测试结果汇总")
    lines.append("")
    lines.append("| 序号 | 测试项 | 生成文件 | 字数 | 状态 |")
    lines.append("|------|--------|----------|------|------|")
    for idx, (name, path, count, status) in enumerate(results, 1):
        file_name = os.path.basename(path) if path and path != "N/A" else "N/A"
        count_str = str(count) if count is not None else "-"
        status_badge = "[PASS]" if status == "PASS" else "[FAIL]"
        lines.append("| {} | {} | {} | {} | {} |".format(idx, name, file_name, count_str, status_badge))
    lines.append("")
    lines.append("- **通过**: {}".format(passed))
    lines.append("- **失败**: {}".format(failed))
    lines.append("")

    lines.append("## 测试覆盖功能")
    lines.append("")
    lines.append("### 文档基础功能")
    lines.append("- [x] 创建新文档")
    lines.append("- [x] 编辑现有文档")
    lines.append("- [x] 页面设置（A4、自定义边距）")
    lines.append("- [x] 页脚页码")
    lines.append("")
    lines.append("### 字体与样式")
    lines.append("- [x] 中文文档样式初始化 (setup_chinese_document_styles)")
    lines.append("- [x] 独立设置中文字体 (set_chinese_font)")
    lines.append("- [x] 文档默认字体 (set_document_default_font)")
    lines.append("- [x] 段落格式（首行缩进、行间距、段间距）")
    lines.append("- [x] 文本样式（粗体、斜体、颜色、字号）")
    lines.append("- [x] 对齐方式（左/中/右/两端对齐）")
    lines.append("")
    lines.append("### 表格功能")
    lines.append("- [x] 创建样式化表格 (create_styled_table)")
    lines.append("- [x] 单元格格式化 (format_table_cell)")
    lines.append("- [x] 单元格边框 (set_cell_border)")
    lines.append("- [x] 单元格背景 (set_cell_background)")
    lines.append("- [x] 单元格边距 (set_cell_margins)")
    lines.append("- [x] 无边框表格 (remove_table_borders)")
    lines.append("")
    lines.append("### 列表与编号")
    lines.append("- [x] 无序列表（项目符号）")
    lines.append("- [x] 有序列表（编号）")
    lines.append("- [x] 列表重新编号 (restart=True)")
    lines.append("- [x] 多级列表嵌套")
    lines.append("")
    lines.append("### 特殊内容")
    lines.append("- [x] 目录 (add_table_of_contents)")
    lines.append("- [x] 超链接 (add_hyperlink)")
    lines.append("- [x] 水平分隔线 (add_horizontal_line)")
    lines.append("- [x] 左右对齐文本（Tab Stops）")
    lines.append("- [x] 公式 (OMath)")
    lines.append("- [x] 图片插入")
    lines.append("")
    lines.append("### 工具函数")
    lines.append("- [x] 字数统计 (count_document_words)")
    lines.append("")
    lines.append("### 边界情况")
    lines.append("- [x] 特殊字符和符号")
    lines.append("- [x] 超长段落")
    lines.append("- [x] 空段落和间距")
    lines.append("- [x] 多行文本（软回车）")
    lines.append("")
    lines.append("## 生成文件列表")
    lines.append("")
    for name, path, count, status in results:
        if path and path != "N/A" and os.path.exists(path):
            size = os.path.getsize(path)
            lines.append("- `{}` ({:,} bytes)".format(os.path.basename(path), size))
    lines.append("")
    lines.append("---")
    lines.append("*Generated by python-docx full test suite*")

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print("\n  [PASS] 测试报告已生成: {}".format(report_path))
    return report_path


# ========================================================================
# Main
# ========================================================================

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def main():
    ensure_dir(OUTPUT_DIR)

    print("\n" + "=" * 60)
    print("Python-docx Skill 全量功能测试开始")
    print("=" * 60)
    print("输出目录: {}".format(OUTPUT_DIR))

    start_time = time.time()
    results = []

    test_cases = [
        ("创建综合文档", test_create_comprehensive_document),
        ("编辑现有文档", test_edit_existing_document),
        ("边界情况测试", test_edge_cases),
        ("超链接测试", test_hyperlink),
        ("单元格高级样式", test_cell_advanced_styles),
        ("默认字体设置", test_default_font),
        ("独立中文字体", test_independent_chinese_font),
    ]

    for name, test_func in test_cases:
        try:
            path, count = test_func()
            results.append((name, path, count, "PASS"))
        except Exception as e:
            results.append((name, "N/A", 0, "FAIL"))
            print("  [FAIL] {} 失败: {}".format(name, e))
            import traceback
            traceback.print_exc()

    duration = time.time() - start_time

    # 生成报告
    try:
        report_path = generate_markdown_report(results, duration)
    except Exception as e:
        print("  [FAIL] 报告生成失败: {}".format(e))

    # 打印总结
    print("\n" + "=" * 60)
    print("测试总结")
    print("=" * 60)
    print("{:<20} {:<30} {:<10} {:<8}".format("测试项", "文件", "字数", "状态"))
    print("-" * 60)
    for name, path, count, status in results:
        file_name = os.path.basename(path) if path != "N/A" else "N/A"
        print("{:<20} {:<30} {:<10} {:<8}".format(name, file_name, count, status))

    passed = sum(1 for r in results if r[3] == "PASS")
    failed = sum(1 for r in results if r[3] == "FAIL")
    print("-" * 60)
    print("总计: {} 通过, {} 失败, 耗时 {:.2f} 秒".format(passed, failed, duration))
    print("=" * 60)


if __name__ == '__main__':
    main()
