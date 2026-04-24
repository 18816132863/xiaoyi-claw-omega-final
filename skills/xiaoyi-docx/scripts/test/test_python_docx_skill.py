#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Python-docx Skill 综合测试脚本
测试内容：
1. 创建新文档 + 编辑现有文档
2. 中英文字体支持
3. 公式 (OMath)
4. 图片插入
5. 表格
6. 列表（有序/无序）
"""

import sys
import os

# 添加 skill 路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 导入工具函数
from docx_utils import (
    setup_chinese_document_styles,
    add_table_of_contents,
    add_page_number_footer,
    count_document_words,
    add_horizontal_line,
    create_styled_table,
    format_table_cell,
    remove_table_borders
)


def add_formula(paragraph, formula_text):
    """添加公式辅助函数"""
    # 创建 oMath 元素
    oMath = OxmlElement('m:oMath')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = formula_text
    r.append(t)
    oMath.append(r)
    paragraph._p.append(oMath)
    return paragraph


def create_test_document():
    """测试1: 创建包含所有功能的新文档"""
    print("=" * 60)
    print("测试1: 创建综合测试文档")
    print("=" * 60)

    doc = Document()

    # 设置中文文档样式
    setup_chinese_document_styles(doc)

    # 设置自定义页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    # ===== 1. 标题 =====
    print("  - 添加标题...")
    title = doc.add_paragraph('Python-docx Skill 综合测试文档', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 2. 目录 =====
    print("  - 添加目录...")
    doc.add_paragraph('目录', style='Heading 1')
    add_table_of_contents(doc)
    doc.add_page_break()

    # ===== 3. 中英文字体测试 =====
    print("  - 测试中英文字体...")
    doc.add_paragraph('第一章 中英文字体测试', style='Heading 1')

    # 中文正文
    p = doc.add_paragraph('这是一段中文正文内容，用于测试 SimSun（宋体）字体显示效果。'
                         '中文文档通常需要设置首行缩进两个字符。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5

    # 英文正文
    p = doc.add_paragraph('This is English text to test font rendering. '
                         'The quick brown fox jumps over the lazy dog. '
                         'Python-docx provides powerful document generation capabilities.',
                         style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 中英混排
    p = doc.add_paragraph('这是一段Mixed Content：中文和English混排测试。'
                         '公式如 E=mc² 和变量 x, y, z 应该正确显示。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 不同样式文本
    doc.add_paragraph('1.1 文本样式测试', style='Heading 2')

    p = doc.add_paragraph()
    p.add_run('这是普通文本 ')
    run_bold = p.add_run('粗体文本 ')
    run_bold.font.bold = True
    run_italic = p.add_run('斜体文本 ')
    run_italic.font.italic = True
    run_color = p.add_run('红色文本 ')
    run_color.font.color.rgb = RGBColor(255, 0, 0)
    run_size = p.add_run('大号文本')
    run_size.font.size = Pt(16)

    # ===== 4. 公式测试 =====
    print("  - 添加公式...")
    doc.add_paragraph('第二章 公式测试', style='Heading 1')

    p = doc.add_paragraph('二次方程求根公式：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 使用 Office Math 添加公式
    p_formula = doc.add_paragraph()
    p_formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_elem = OxmlElement('m:oMath')
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = 'x = '
    r1.append(t1)

    # 构建分数
    frac = OxmlElement('m:f')
    num = OxmlElement('m:num')
    r_num = OxmlElement('w:r')
    t_num = OxmlElement('w:t')
    t_num.text = '-b ± √(b²-4ac)'
    r_num.append(t_num)
    num.append(r_num)

    den = OxmlElement('m:den')
    r_den = OxmlElement('w:r')
    t_den = OxmlElement('w:t')
    t_den.text = '2a'
    r_den.append(t_den)
    den.append(r_den)

    frac.append(num)
    frac.append(den)
    math_elem.append(r1)
    math_elem.append(frac)
    p_formula._p.append(math_elem)

    p = doc.add_paragraph('积分公式示例：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p_formula2 = doc.add_paragraph()
    p_formula2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_elem2 = OxmlElement('m:oMath')
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = '∫[a,b] f(x)dx = F(b) - F(a)'
    r2.append(t2)
    math_elem2.append(r2)
    p_formula2._p.append(math_elem2)

    # ===== 5. 图片测试 =====
    print("  - 添加图片...")
    doc.add_paragraph('第三章 图片测试', style='Heading 1')

    p = doc.add_paragraph('以下是图片插入测试：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 创建一个简单的测试图片 (使用 ASCII art 作为文本代替，或尝试找现有图片)
    # 由于没有实际图片文件，我们创建一个临时的
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (400, 200), color='lightblue')
        draw = ImageDraw.Draw(img)
        draw.text((50, 80), "Test Image\nPython-docx", fill='black')
        img_path = 'test_image_temp.png'
        img.save(img_path)

        doc.add_picture(img_path, width=Inches(4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 清理临时文件
        os.remove(img_path)
    except ImportError:
        p = doc.add_paragraph('[PIL 未安装，跳过图片测试]', style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 6. 表格测试 =====
    print("  - 添加表格...")
    doc.add_paragraph('第四章 表格测试', style='Heading 1')

    p = doc.add_paragraph('以下是 styled table 测试：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 创建样式化表格
    table = create_styled_table(doc, rows=4, cols=3, col_widths=[2.0, 2.5, 3.0])

    # 表头
    format_table_cell(table.rows[0].cells[0], "功能模块", alignment='center')
    format_table_cell(table.rows[0].cells[1], "支持状态", alignment='center')
    format_table_cell(table.rows[0].cells[2], "说明", alignment='center')

    # 数据行
    table.rows[1].cells[0].text = "中英文字体"
    format_table_cell(table.rows[1].cells[1], "✓ 支持", alignment='center')
    table.rows[1].cells[2].text = "SimSun/SimHei 等中文字体"

    table.rows[2].cells[0].text = "公式编辑"
    format_table_cell(table.rows[2].cells[1], "✓ 支持", alignment='center')
    table.rows[2].cells[2].text = "OMath 公式支持"

    table.rows[3].cells[0].text = "图片插入"
    format_table_cell(table.rows[3].cells[1], "✓ 支持", alignment='center')
    table.rows[3].cells[2].text = "支持 PNG/JPG 等格式"

    # 无边框表格示例
    doc.add_paragraph('')
    p = doc.add_paragraph('无边框表格（用于布局）：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    layout_table = doc.add_table(rows=2, cols=2)
    layout_table.rows[0].cells[0].text = "左侧内容"
    layout_table.rows[0].cells[1].text = "右侧内容"
    layout_table.rows[1].cells[0].text = "标签："
    layout_table.rows[1].cells[1].text = "数值内容"
    remove_table_borders(layout_table)

    # ===== 7. 列表测试 =====
    print("  - 添加列表...")
    doc.add_paragraph('第五章 列表测试', style='Heading 1')

    # 无序列表
    doc.add_paragraph('5.1 无序列表（项目符号）', style='Heading 2')
    doc.add_paragraph('第一项：Python-docx 基础功能', style='List Bullet')
    doc.add_paragraph('第二项：中文文档支持', style='List Bullet')
    doc.add_paragraph('第三项：样式和格式设置', style='List Bullet')
    doc.add_paragraph('嵌套项示例', style='List Bullet 2')
    doc.add_paragraph('更深嵌套', style='List Bullet 3')

    # 有序列表
    doc.add_paragraph('5.2 有序列表（编号）', style='Heading 2')
    doc.add_paragraph('安装 python-docx 库', style='List Number', restart=True)
    doc.add_paragraph('导入 Document 类', style='List Number')
    doc.add_paragraph('创建文档对象', style='List Number')
    doc.add_paragraph('添加内容和样式', style='List Number')
    doc.add_paragraph('保存文档', style='List Number')

    # 新列表（重置编号）
    doc.add_paragraph('')
    doc.add_paragraph('重新开始编号：', style='Normal')
    doc.add_paragraph('第一步：分析问题', style='List Number', restart=True)
    doc.add_paragraph('第二步：设计方案', style='List Number')
    doc.add_paragraph('第三步：实施开发', style='List Number')

    # ===== 8. 水平分隔线测试 =====
    print("  - 添加水平分隔线...")
    doc.add_paragraph('第六章 水平分隔线', style='Heading 1')

    p = doc.add_paragraph('以下是单线分隔：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    line1 = doc.add_paragraph()
    add_horizontal_line(line1)

    p = doc.add_paragraph('以下是双线分隔（红色）：', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)
    line2 = doc.add_paragraph()
    add_horizontal_line(line2, line_style='double', line_size=12, color='FF0000')

    # ===== 9. 左右对齐文本 =====
    print("  - 添加左右对齐文本...")
    doc.add_paragraph('第七章 特殊排版', style='Heading 1')

    p = doc.add_paragraph('左右对齐文本示例：', style='Heading 2')

    # 计算页面宽度
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    p_align = doc.add_paragraph()
    p_align.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    p_align.add_run('左侧文本（左对齐）')
    run_tab = p_align.add_run('\t')
    run_tab.underline = True
    p_align.add_run('右侧文本（右对齐）')

    # ===== 10. 统计字数 =====
    print("  - 统计字数...")
    word_count = count_document_words(doc)
    doc.add_paragraph('第八章 文档统计', style='Heading 1')

    p = doc.add_paragraph(f'本文档总字数：{word_count} 字', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # ===== 11. 页脚页码 =====
    print("  - 添加页码...")
    add_page_number_footer(section, alignment='center')

    # 保存文档
    output_path = 'test_output_create.docx'
    doc.save(output_path)
    print(f"  ✓ 文档已保存: {output_path}")

    return output_path, word_count


def edit_existing_document():
    """测试2: 编辑现有文档"""
    print("\n" + "=" * 60)
    print("测试2: 编辑现有文档")
    print("=" * 60)

    # 先创建一个基础文档
    doc = Document()
    setup_chinese_document_styles(doc)

    doc.add_paragraph('原始文档标题', style='Title')
    doc.add_paragraph('第一章 原始内容', style='Heading 1')
    p = doc.add_paragraph('这是原始文档的第一段内容。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 保存原始文档
    original_path = 'test_original.docx'
    doc.save(original_path)
    print(f"  - 创建原始文档: {original_path}")

    # 重新打开并编辑
    doc = Document(original_path)
    print(f"  - 重新打开文档进行编辑")

    # 添加新章节
    doc.add_paragraph('第二章 新增内容', style='Heading 1')
    p = doc.add_paragraph('这是在编辑时添加的新段落。测试文档追加功能。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 添加表格
    table = create_styled_table(doc, rows=3, cols=2, col_widths=[3.0, 4.0])
    format_table_cell(table.rows[0].cells[0], "属性", alignment='center')
    format_table_cell(table.rows[0].cells[1], "值", alignment='center')
    table.rows[1].cells[0].text = "文档类型"
    table.rows[1].cells[1].text = "测试文档"
    table.rows[2].cells[0].text = "编辑次数"
    table.rows[2].cells[1].text = "1"

    # 添加列表
    doc.add_paragraph('编辑步骤：', style='Heading 2')
    doc.add_paragraph('打开现有文档', style='List Number', restart=True)
    doc.add_paragraph('添加新内容', style='List Number')
    doc.add_paragraph('保存修改', style='List Number')

    # 添加页码
    section = doc.sections[0]
    add_page_number_footer(section, alignment='right')

    # 保存编辑后的文档
    output_path = 'test_output_edit.docx'
    doc.save(output_path)
    print(f"  ✓ 编辑后的文档已保存: {output_path}")

    # 统计字数
    word_count = count_document_words(output_path)
    print(f"  - 编辑后文档字数: {word_count}")

    # 清理原始文件
    os.remove(original_path)

    return output_path, word_count


def test_edge_cases():
    """测试3: 边界情况和特殊字符"""
    print("\n" + "=" * 60)
    print("测试3: 边界情况和特殊字符")
    print("=" * 60)

    doc = Document()
    setup_chinese_document_styles(doc)

    doc.add_paragraph('边界情况测试', style='Title')

    # 特殊字符
    doc.add_paragraph('特殊字符测试', style='Heading 1')
    p = doc.add_paragraph('特殊符号：© ® ™ ° ℃ ℉ § № ★ ☆ ▲ ▼ ◆ ◇ ● ○ ', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p = doc.add_paragraph('数学符号：± × ÷ ∞ √ ∑ ∏ ∫ ∮ ≈ ≠ ≤ ≥ ', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    p = doc.add_paragraph('货币符号：￥ $ € £ ¥ ₹ ', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 超长段落
    doc.add_paragraph('长段落测试', style='Heading 1')
    long_text = '这是一段很长的文本内容，用于测试段落处理和分页功能。' * 20
    p = doc.add_paragraph(long_text, style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 空段落和间距
    doc.add_paragraph('间距测试', style='Heading 1')
    p = doc.add_paragraph('这是一个有下方间距的段落。', style='Normal')
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph('上方段落有24pt间距。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 多行文本（使用软回车）
    doc.add_paragraph('多行文本测试', style='Heading 1')
    p = doc.add_paragraph('第一行\n第二行（软回车）\n第三行（软回车）', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24)

    # 对齐方式测试
    doc.add_paragraph('对齐方式测试', style='Heading 1')
    p = doc.add_paragraph('左对齐文本（默认）')
    p = doc.add_paragraph('居中文本')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('右对齐文本')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('两端对齐文本，用于测试文本两端对齐效果。')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output_path = 'test_output_edge_cases.docx'
    doc.save(output_path)
    print(f"  ✓ 边界测试文档已保存: {output_path}")

    word_count = count_document_words(output_path)
    print(f"  - 文档字数: {word_count}")

    return output_path, word_count


def main():
    """主函数：运行所有测试"""
    print("\n" + "=" * 60)
    print("Python-docx Skill 综合测试开始")
    print("=" * 60)

    results = []

    # 测试1: 创建文档
    try:
        path1, count1 = create_test_document()
        results.append(("创建文档测试", path1, count1, "✓ 通过"))
    except Exception as e:
        results.append(("创建文档测试", "N/A", 0, f"✗ 失败: {e}"))

    # 测试2: 编辑文档
    try:
        path2, count2 = edit_existing_document()
        results.append(("编辑文档测试", path2, count2, "✓ 通过"))
    except Exception as e:
        results.append(("编辑文档测试", "N/A", 0, f"✗ 失败: {e}"))

    # 测试3: 边界情况
    try:
        path3, count3 = test_edge_cases()
        results.append(("边界情况测试", path3, count3, "✓ 通过"))
    except Exception as e:
        results.append(("边界情况测试", "N/A", 0, f"✗ 失败: {e}"))

    # 打印测试总结
    print("\n" + "=" * 60)
    print("测试总结")
    print("=" * 60)
    print(f"{'测试项':<20} {'文件':<25} {'字数':<10} {'状态'}")
    print("-" * 60)
    for name, path, count, status in results:
        print(f"{name:<20} {path:<25} {count:<10} {status}")

    print("\n" + "=" * 60)
    print("测试完成！生成的文件：")
    for name, path, _, _ in results:
        if path != "N/A":
            print(f"  - {path}")
    print("=" * 60)


if __name__ == '__main__':
    main()
